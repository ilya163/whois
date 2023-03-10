from domains.my_additional import *
from docx import Document
from docx.shared import Mm, Pt

class SaveWord:

    @classmethod
    def save(cls, name_domain, response):
        try:
            if not os.path.exists(PATH_WORD_FILE):
                doc = Document()
            else:
                doc = Document(PATH_WORD_FILE)
                doc.add_page_break()

            from docx.enum.style import WD_STYLE_TYPE
            table = doc.add_table(rows=0, cols=2, style="Light Grid")

            results = cls.formed_data(response)

            for row in results:
                # добавляем строку с ячейками к объекту таблицы
                cells = table.add_row().cells
                for i, item in enumerate(row, 0):
                    if row[i-1] == "Скриншот":
                        PATH_SCREENSHOT = os.path.join(PATH_RESULTS_DIR, name_domain, "screenshot.png")
                        if os.path.exists(PATH_SCREENSHOT):
                            p = cells[1].add_paragraph()
                            run = p.add_run()
                            run.add_picture(PATH_SCREENSHOT, width=Mm(100), height=Mm(60))
                            continue
                    else:
                        cells[i].text = str(item)

            doc.save(PATH_WORD_FILE)
        except Exception as err:
            logger.error(f"Ошибка при сохранении в Word файл {str(err)}")


    @staticmethod
    def formed_data(response):
        try:
            if "main" in response:
                main = response["main"]
                results = [["Анализируемый домен", main.get("domain")],
                           ["IP-адрес", main.get("ip")],
                           ["Владелец", main.get("owner")],
                           ["Доступность сайта", response.get('state'),],
                           ["Регистратор", main.get("registrar")],
                           ["Создан", main.get("create_time")],
                           ["На последнем ip-адресе", main.get("on_last_ip_address")],
                           ["Оплачен до", main.get("expires")],
                           ["Статус", main.get("status")],
                           ["Наличие архивных версий сайта", main.get("archive")],
                           ["Список DNS-серверов", field_ns_server(main.get("ns"))],
                           ["Скриншот", "Отсутствует"],
                           ["Наличие архивных версий сайта", main.get("archive")],
                           ["Связанные адреса эл.почты", linked_email_save(main.get("linked_email"))],
                           ['Поддомены', worker_write_list(main.get("subdomains"))],
                           ['Данные указанные на сайте', dict_to_str(main.get("metadata"))],
                           ]

                if "contact" in main:
                    if len(main["contact"]) > 0:
                        result = worker_write_list(main["contact"],
                                                          header='Контакты')
                        results.insert(1, result)

                if "ip_history" in main:
                    if len(main["ip_history"]) > 0:
                        result = worker_dynamic_type(main["ip_history"],
                                                          header='IP-адреса на которых размещался домен')

                        results.append(result)
                if "other_domain" in main:
                    if len(main["other_domain"]) > 0:
                        result = worker_dynamic_type(main["other_domain"],
                                                          header='Прочие домены на сервере', typeList=True)

                        results.append(result)

                if "rw_domain" in main:
                    if len(main["rw_domain"]) > 0:
                        result = worker_dynamic_type(main["rw_domain"],
                                                          header='Прочие домены принадлежащие владельцу', typeList=True)
                        results.append(result)
                return results

        except Exception as err:
            logger.error(f"Ошибка при запуске обработчика Word: {str(err)}")
